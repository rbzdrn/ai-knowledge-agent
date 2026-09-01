#!/usr/bin/env python3
"""Generate the interview preparation Word document for AI Knowledge Base Agent."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Style Configuration ──────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Heading styles
for level in range(1, 5):
    heading_style = doc.styles[f'Heading {level}']
    hf = heading_style.font
    hf.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hf.color.rgb = RGBColor(0x1a, 0x1d, 0x23)
    if level == 1:
        hf.size = Pt(22)
    elif level == 2:
        hf.size = Pt(16)
    elif level == 3:
        hf.size = Pt(13)
    elif level == 4:
        hf.size = Pt(11)

# Create a code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(8.5)
code_font.color.rgb = RGBColor(0x1a, 0x1d, 0x23)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.left_indent = Cm(0.5)

def add_paragraph(text, bold=False, style_name=None):
    """Add a normal paragraph."""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    return p

def add_code(code_text):
    """Add a code block with grey background shading."""
    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph(style='CodeBlock')
        # Add grey background
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F5F5F5')
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)

def add_bullet(text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p

def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacing
    return table

def add_separator():
    """Add a horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI Knowledge Base Agent')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('面试准备资料 — 项目流程、代码详解与面试问答')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph()
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('技术栈：FastAPI + ChromaDB + Claude + BM25 + Cytoscape.js\n').font.size = Pt(10)
meta.add_run('涵盖：架构设计 · 模块源码解析 · 30道面试题详解').font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目概述',
    '二、系统架构与数据流',
    '  2.1 架构全景图',
    '  2.2 文档摄入流（离线）',
    '  2.3 问答流（在线）',
    '三、模块代码详解',
    '  3.1 run.py — 启动入口',
    '  3.2 config.py — 配置管理',
    '  3.3 embeddings.py — 嵌入模型',
    '  3.4 document_loader/loader.py — 文档解析',
    '  3.5 chunker/chunker.py — CJK感知文本切分',
    '  3.6 vectordb/store.py — ChromaDB向量存储',
    '  3.7 retriever/hybrid.py — 混合检索 + BM25 + RRF',
    '  3.8 llm.py — 多后端LLM抽象',
    '  3.9 agent/agent.py — 知识问答Agent',
    '  3.10 agent/prompts.py — 提示词设计',
    '  3.11 kg/builder.py — 知识图谱',
    '  3.12 main.py — FastAPI路由',
    '  3.13 前端 app.js — SSE流式处理',
    '四、面试问题与答案（30题）',
    '  4.1 RAG基础（Q1-Q3）',
    '  4.2 嵌入与向量检索（Q4-Q7）',
    '  4.3 BM25与混合检索（Q8-Q10）',
    '  4.4 LLM相关（Q11-Q15）',
    '  4.5 知识图谱（Q16-Q17）',
    '  4.6 工程实践（Q18-Q22）',
    '  4.7 系统设计（Q23-Q24）',
    '  4.8 前端（Q25-Q26）',
    '  4.9 综合问题（Q27-Q30）',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run(item).font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  PART 1: PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════
doc.add_heading('一、项目概述', level=1)

add_paragraph(
    'AI Knowledge Base Agent 是一个基于 RAG（Retrieval-Augmented Generation，检索增强生成）'
    '的智能知识库系统。用户可以上传多种格式的文档（PDF、DOCX、Markdown、网页等），系统自动'
    '完成文档解析、文本切分、向量化存储，并通过混合检索引擎和 AI Agent 提供带原文溯源的智能问答。'
)

doc.add_heading('1.1 核心能力', level=2)
add_table(
    ['能力', '描述'],
    [
        ['多格式文档摄入', 'PDF / DOCX / Markdown / TXT / 网页，自动解析并索引'],
        ['混合检索', 'Dense（向量相似度）+ Sparse（BM25 关键词）+ RRF 融合排序'],
        ['AI 智能问答', 'LLM 驱动，基于检索到的上下文生成答案，附带来源引用'],
        ['多轮对话', '会话级别的对话历史管理，支持上下文连续问答'],
        ['知识图谱', '自动实体抽取 + 文档-实体关系图可视化'],
        ['流式响应', 'SSE（Server-Sent Events）实现 token 级别实时输出'],
    ]
)

doc.add_heading('1.2 技术栈', level=2)
add_table(
    ['层级', '技术选型', '选型理由'],
    [
        ['Web 框架', 'FastAPI + Uvicorn', '原生异步支持，SSE 流式响应，自动生成 API 文档'],
        ['向量数据库', 'ChromaDB（嵌入式）', '零配置，零运维，SQLite 持久化，单机百万级 chunk 无压力'],
        ['嵌入模型', 'all-MiniLM-L6-v2', '本地运行，80MB 体积，384 维向量，CPU 推理毫秒级'],
        ['LLM', 'Anthropic Claude / Ollama / DeepSeek / OpenAI', '多后端自动检测，从免费到付费灵活切换'],
        ['关键词检索', '自研 BM25', '轻量级，CJK 分词感知，与 Dense 检索互补'],
        ['知识图谱', 'NetworkX + Cytoscape.js', '轻量图计算 + 前端交互式可视化'],
        ['前端', '原生 HTML/CSS/JS', '零构建工具，单文件部署，Cytoscape CDN 引入'],
    ]
)

doc.add_heading('1.3 项目结构', level=2)
add_code('''ai-knowledge-agent/
├── run.py                       # 启动脚本（uvicorn）
├── README.md                    # 项目说明
├── docs/                        # 文档
│   ├── ARCHITECTURE.md          # 架构设计
│   ├── API.md                   # API 参考
│   └── DEVELOPMENT.md           # 开发指南
├── backend/
│   ├── main.py                  # FastAPI 应用 + 全部 API 路由
│   ├── config.py                # 全局配置（dataclass 单例）
│   ├── embeddings.py            # 嵌入模型封装（懒加载）
│   ├── llm.py                   # 多后端 LLM 抽象层
│   ├── requirements.txt         # 依赖清单
│   ├── document_loader/
│   │   └── loader.py            # 多格式文档解析（PDF/DOCX/MD/TXT/Web）
│   ├── chunker/
│   │   └── chunker.py           # CJK 感知递归文本切分器
│   ├── vectordb/
│   │   └── store.py             # ChromaDB 封装（CRUD + 搜索）
│   ├── retriever/
│   │   └── hybrid.py            # BM25 实现 + HybridRetriever + RRF 融合
│   ├── agent/
│   │   ├── agent.py             # 知识问答 Agent（检索+生成+会话管理）
│   │   └── prompts.py           # System Prompt（含引用规则）
│   └── kg/
│       └── builder.py           # 实体抽取 + NetworkX 图构建 + Cytoscape 导出
├── frontend/
│   ├── index.html               # SPA 入口（三栏布局）
│   ├── style.css                # 响应式样式（CSS 变量 + 暗色侧栏）
│   └── app.js                   # 聊天 / SSE 流式 / 文件上传 / 图谱渲染
└── data/                        # 运行时自动创建
    ├── uploads/                  # 上传的文档
    └── vectordb/                 # ChromaDB 持久化存储（chroma.sqlite3）''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  PART 2: ARCHITECTURE
# ═══════════════════════════════════════════════════════════
doc.add_heading('二、系统架构与数据流', level=1)

doc.add_heading('2.1 架构全景图', level=2)
add_code('''┌─────────────────────────────────────────────────────────────────────┐
│                        Frontend (SPA)                                │
│  ┌──────────┐  ┌──────────────┐  ┌──────────────┐  ┌────────────┐  │
│  │ 文档上传  │  │  聊天界面     │  │  来源面板     │  │  知识图谱   │  │
│  │ (拖拽/URL)│  │  (SSE 流式)  │  │  (溯源引用)   │  │  (Cytoscape)│  │
│  └────┬─────┘  └──────┬───────┘  └──────┬───────┘  └─────┬──────┘  │
└───────┼───────────────┼─────────────────┼───────────────┼──────────┘
        │               │                 │               │
┌───────┴───────────────┴─────────────────┴───────────────┴──────────┐
│                     FastAPI Server                                  │
│                                                                     │
│  ┌──────────────────────────────────────────────────────────────┐  │
│  │                    REST API Layer                              │  │
│  │  /upload  /upload-url  /documents  /chat  /chat/stream        │  │
│  │  /search  /knowledge-graph  /conversation/reset  /health      │  │
│  └──────────────────────────────────────────────────────────────┘  │
│                              │                                      │
│  ┌───────────────────────────┴──────────────────────────────────┐  │
│  │                     Core Pipeline                             │  │
│  │                                                               │  │
│  │   ┌─────────┐   ┌──────────┐   ┌──────────┐   ┌──────────┐   │  │
│  │   │ Loader  │──▶│ Chunker  │──▶│ Embedder │──▶│VectorStore│  │  │
│  │   │ PDF/DOCX│   │ CJK-aware│   │ MiniLM   │   │ ChromaDB  │   │  │
│  │   │ MD/Web  │   │ splitter │   │ L6-v2    │   │ persistent│   │  │
│  │   └─────────┘   └──────────┘   └──────────┘   └─────┬─────┘   │  │
│  │                                                      │         │  │
│  │   ┌──────────────────────────────────────────────────┘         │  │
│  │   ▼                                                            │  │
│  │   ┌──────────────┐   ┌──────────────┐   ┌──────────────┐      │  │
│  │   │ BM25 Index   │   │HybridRetriever│   │ Knowledge    │      │  │
│  │   │ (in-memory)  │◀──│ Dense+BM25    │   │ Agent        │      │  │
│  │   │              │   │ + RRF fusion  │──▶│ (LLM)        │      │  │
│  │   └──────────────┘   └──────────────┘   └──────┬───────┘      │  │
│  │                                                  │              │  │
│  │   ┌──────────────┐                              ▼              │  │
│  │   │ Knowledge    │                     Streaming Response      │  │
│  │   │ Graph Builder│                     (SSE token-by-token)    │  │
│  │   │ NetworkX     │                                              │  │
│  │   └──────────────┘                                              │  │
│  └─────────────────────────────────────────────────────────────────┘  │
└─────────────────────────────────────────────────────────────────────┘''')

doc.add_heading('2.2 文档摄入流（离线）', level=2)
add_paragraph('文档从上传到可检索的完整处理链路：')
add_code('''文件上传 → 保存到 data/uploads/（UUID前缀防冲突）
         → DocumentLoader 解析文本（PDF→pypdf, DOCX→python-docx+XML fallback）
         → TextChunker 切分为 chunk（CJK感知，500字/chunk，50字重叠）
         → Embedder 批量向量化（384维，L2归一化）
         → ChromaDB 持久化存储（HNSW索引，cosine距离度量）
         → 全量重建 BM25 索引（更新 DF / avgdl / corpus）''')

doc.add_heading('2.3 问答流（在线）', level=2)
add_paragraph('用户提问后的实时处理链路：')
add_code('''用户提问 → Embedder 将问题向量化（与文档使用同一嵌入空间）
         → Dense 检索: ChromaDB cosine相似度查询 top_k=10
         → Sparse 检索: BM25 关键词匹配 top_k=10
         → RRF 融合: 两个排序列表按倒数排名加权合并，k=60
         → 取 top_k=8 个最终 chunk
         → 组装 System Prompt: SYSTEM_PROMPT.format(context=检索上下文)
         → 获取会话历史 + 拼接用户消息
         → 调用 LLM（同步 / 流式 SSE）
         → 提取来源信息（按文档分组，含文件名/chunk索引/相关性分数）
         → 更新会话历史（user消息 + assistant消息）''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  PART 3: MODULE CODE DETAILS
# ═══════════════════════════════════════════════════════════
doc.add_heading('三、模块代码详解', level=1)

# 3.1
doc.add_heading('3.1 run.py — 启动入口', level=2)
add_paragraph('作用：使用 uvicorn 启动 FastAPI 应用，配置热重载和 host/port。')
add_code('''#!/usr/bin/env python3
"""Start the AI Knowledge Base Agent server."""
import sys
from pathlib import Path

# 关键: 将项目根目录加入 sys.path，确保 backend 包可被导入
sys.path.insert(0, str(Path(__file__).resolve().parent))

import uvicorn

if __name__ == "__main__":
    print("Starting AI Knowledge Base Agent...")
    uvicorn.run(
        "backend.main:app",    # FastAPI 应用路径
        host="0.0.0.0",        # 监听所有网络接口
        port=8000,
        reload=True,            # 开发模式热重载
        log_level="info",
    )''')
add_paragraph('关键点：')
add_bullet('sys.path.insert(0, ...) 确保 backend 包可以被正确导入')
add_bullet('reload=True 开发模式下代码变更自动重启')
add_bullet('host="0.0.0.0" 允许局域网内其他设备访问')

# 3.2
doc.add_heading('3.2 config.py — 配置管理', level=2)
add_paragraph('使用 Python dataclass 实现全局配置单例，所有模块通过 from .config import config 引用。')
add_code('''@dataclass
class Config:
    # ── 路径 ──
    data_dir: Path = DATA_DIR
    vectordb_dir: Path = VECTORDB_DIR
    upload_dir: Path = UPLOAD_DIR

    # ── 嵌入模型 ──
    embedding_model: str = "all-MiniLM-L6-v2"   # 80MB 本地模型
    embedding_dim: int = 384                     # 向量维度
    embedding_device: str = "cpu"                # "cpu" 或 "cuda"

    # ── 文本切分 ──
    chunk_size: int = 500        # 每个 chunk 的目标字符数
    chunk_overlap: int = 50      # 相邻 chunk 重叠字符数（10%）

    # ── 检索 ──
    dense_top_k: int = 10        # Dense 检索候选数
    sparse_top_k: int = 10       # BM25 检索候选数
    hybrid_top_k: int = 8        # RRF 融合后最终返回数
    use_reranker: bool = False   # Cross-encoder 重排序开关

    # ── Agent ──
    llm_model: str = "claude-sonnet-4-6"
    llm_max_tokens: int = 2048
    conversation_ttl: int = 3600  # 会话过期时间（秒）

    # ── 知识图谱 ──
    kg_max_entities_per_chunk: int = 5

    def __post_init__(self):
        """自动创建数据目录"""
        self.data_dir.mkdir(parents=True, exist_ok=True)
        self.vectordb_dir.mkdir(parents=True, exist_ok=True)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    @property
    def anthropic_api_key(self) -> str:
        return os.environ.get("ANTHROPIC_API_KEY", "")

# 模块级单例：整个应用共用这一个实例
config = Config()''')
add_paragraph('设计模式：dataclass + 模块级单例 = config。__post_init__ 自动创建数据目录，属性方法从环境变量动态读取 API Key。')

# 3.3
doc.add_heading('3.3 embeddings.py — 嵌入模型', level=2)
add_paragraph('封装 sentence-transformers 模型，实现懒加载和容错策略。')
add_code('''class Embedder:
    """Generate embeddings using local sentence-transformers model."""

    def __init__(self, model_name=None, device=None):
        self.model_name = model_name or config.embedding_model
        self.device = device or config.embedding_device
        self._model = None  # ← 关键：懒加载，不立即加载模型

    @property
    def model(self):
        if self._model is None:
            from sentence_transformers import SentenceTransformer
            kwargs = {"device": self.device}
            # 策略1: 先尝试从本地缓存加载（离线模式）
            try:
                self._model = SentenceTransformer(
                    self.model_name, local_files_only=True, **kwargs
                )
            except Exception:
                # 策略2: 在线下载，自动设置国内镜像
                if not os.environ.get("HF_ENDPOINT"):
                    os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
                print(f"Downloading embedding model (one-time, ~80MB)...")
                self._model = SentenceTransformer(self.model_name, **kwargs)
        return self._model

    def embed(self, text: str) -> List[float]:
        """单条文本向量化，L2归一化"""
        return self.model.encode(text, normalize_embeddings=True).tolist()

    def embed_batch(self, texts: List[str]) -> List[List[float]]:
        """批量文本向量化，提高吞吐量"""
        embeddings = self.model.encode(
            texts, normalize_embeddings=True, show_progress_bar=False
        )
        return embeddings.tolist()

    @property
    def dim(self) -> int:
        """获取向量维度"""
        return self.model.get_sentence_embedding_dimension()''')
add_paragraph('关键设计：')
add_bullet('懒加载：首次调用 .model 属性才加载，模块导入不阻塞启动')
add_bullet('双层容错：先尝试 local_files_only=True（离线/已缓存），失败再在线下载')
add_bullet('国内镜像：自动设置 HF_ENDPOINT=https://hf-mirror.com')
add_bullet('L2 归一化：normalize_embeddings=True 确保余弦相似度计算正确')
add_bullet('批量编码：embed_batch() 单次调用处理多个文本，比循环调用 embed() 快 3-5 倍')

# 3.4
doc.add_heading('3.4 document_loader/loader.py — 文档解析', level=2)
add_paragraph('支持 PDF、DOCX、Markdown、TXT、网页 5 种格式，输出统一的 Document 数据类。')

add_paragraph('核心数据结构：', bold=True)
add_code('''@dataclass
class Document:
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:8])
    filename: str = ""
    content: str = ""
    metadata: dict = field(default_factory=dict)
    chunk_ids: List[str] = field(default_factory=list)''')

add_paragraph('格式分发逻辑：', bold=True)
add_code('''@staticmethod
def load(file_path: str | Path) -> Document:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        content = DocumentLoader._load_pdf(path)      # pypdf.PdfReader
    elif suffix in (".docx", ".doc"):
        content = DocumentLoader._load_docx(path)     # python-docx + XML fallback
    elif suffix in (".md", ".markdown"):
        content = DocumentLoader._load_markdown(path)  # UTF-8 读取
    elif suffix in (".txt", ".text", ".csv", ".json", ".xml", ".html", ".htm"):
        content = DocumentLoader._load_text(path)      # UTF-8 读取
    else:
        raise ValueError(f"Unsupported file format: {suffix}")

    content = DocumentLoader._clean_text(content)
    return Document(
        filename=path.name,
        content=content,
        metadata={"source": str(path), "file_type": suffix, "char_count": len(content)},
    )''')

add_paragraph('DOCX 双重回退机制（解决 WPS/老版本兼容问题）：', bold=True)
add_code('''@staticmethod
def _load_docx(path: Path) -> str:
    # 方案1: python-docx 标准解析
    doc = DocxDocument(str(path))
    text = "\\n".join(p.text for p in doc.paragraphs if p.text.strip())
    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text:
                text += "\\n" + row_text

    # 方案2: 如果标准解析结果为空，直接从 ZIP 内 XML 提取
    if not text.strip():
        text = DocumentLoader._extract_docx_xml(path)
    return text

@staticmethod
def _extract_docx_xml(path: Path) -> str:
    """用正则从 docx 内部 XML 提取 <w:t> 标签文本"""
    import zipfile
    paragraphs = []
    with zipfile.ZipFile(str(path)) as zf:
        for name in zf.namelist():
            if name.endswith('.xml') and 'document' in name.lower():
                xml = zf.read(name).decode('utf-8', errors='ignore')
                # 按段落分组提取
                for para_match in re.finditer(r'<w:p[ >](.*?)</w:p>', xml, re.DOTALL):
                    texts = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', para_match.group(1))
                    line = ''.join(texts).strip()
                    if line:
                        paragraphs.append(line)
    return "\\n".join(paragraphs)''')

add_paragraph('网页解析：', bold=True)
add_code('''@staticmethod
def load_url(url: str) -> Document:
    headers = {"User-Agent": "Mozilla/5.0 (compatible; KnowledgeBot/1.0)"}
    resp = requests.get(url, headers=headers, timeout=30)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    # 去掉非正文标签
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\\n")
    text = DocumentLoader._clean_text(text)
    return Document(
        filename=soup.title.string.strip() if soup.title else url,
        content=text,
        metadata={"source": url, "file_type": "web", "char_count": len(text)},
    )''')

add_paragraph('文本清洗：', bold=True)
add_code('''@staticmethod
def _clean_text(text: str) -> str:
    text = re.sub(r'\\n{3,}', '\\n\\n', text)    # 合并多余空行（3+ → 2）
    text = re.sub(r'[ \\t]{3,}', '  ', text)      # 合并多余空格/制表符
    text = re.sub(r'\\x00', '', text)              # 移除空字符（PDF编码问题）
    return text.strip()''')

# 3.5
doc.add_heading('3.5 chunker/chunker.py — CJK 感知文本切分', level=2)
add_paragraph('这是项目核心创新之一。LangChain 默认分隔符不处理中文标点（。；！？），导致中文文档切分质量差。自研 Chunker 解决了该问题。')

add_code('''class TextChunker:
    """Split text into overlapping chunks for embedding and retrieval."""

    def __init__(self, chunk_size=500, chunk_overlap=50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        # 分隔符优先级：从粗到细
        self._separators = [
            "\\n\\n", "\\n",           # 段落级
            "。", ". ", ".",           # 句子级（中英文句号）
            "；", ";",                 # 分句级
            "！", "!", "？", "?",      # 感叹/疑问
            " ", ""                    # 最终回退：按字符
        ]

    def split(self, text: str) -> List[str]:
        return self._split_recursive(text, self._separators)

    def _split_recursive(self, text, separators):
        """递归切分核心算法"""
        chunks = []
        separator = separators[0]
        next_separators = separators[1:] if len(separators) > 1 else [""]

        if not separator:
            # 最终回退：固定宽度切分
            return self._split_by_length(text)

        splits = text.split(separator)
        current = ""
        for part in splits:
            candidate = current + (separator if current else "") + part
            if len(candidate) <= self.chunk_size:
                current = candidate          # 贪心拼接
            else:
                if current:
                    chunks.append(current.strip())
                if len(part) > self.chunk_size:
                    # 超长片段递归使用下一级（更细）分隔符
                    sub_chunks = self._split_recursive(part, next_separators)
                    # 与前一个 chunk 重叠
                    if chunks and self.chunk_overlap > 0:
                        overlap_text = chunks[-1][-self.chunk_overlap:]
                        if sub_chunks:
                            sub_chunks[0] = overlap_text + sub_chunks[0]
                    chunks.extend(sub_chunks)
                    current = ""
                else:
                    current = part

        if current.strip():
            chunks.append(current.strip())

        # 合并过短 chunk（< 50 字符）
        return self._merge_short(chunks)

    def _split_by_length(self, text):
        """最终回退：按固定 chunk_size 切分 + overlap"""
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunks.append(text[start:end].strip())
            start = end - self.chunk_overlap
        return [c for c in chunks if c]

    def _merge_short(self, chunks, min_len=50):
        """合并过短的 chunk 到前一个 chunk"""
        merged = []
        for chunk in chunks:
            if len(chunk) < min_len and merged:
                merged[-1] = merged[-1] + " " + chunk  # 合并
            else:
                merged.append(chunk)
        return merged''')

add_paragraph('算法要点：')
add_bullet('分隔符优先级从粗到细：段落 → 句子 → 分句 → 词 → 字符')
add_bullet('贪心拼接：只要拼接后不超过 chunk_size 就一直拼接')
add_bullet('递归回退：超长片段使用更细粒度的分隔符继续切分')
add_bullet('Overlap：相邻 chunk 保留 50 字符重叠，防止关键信息落在边界')
add_bullet('短 chunk 合并：< 50 字符的片段合并到前一个，避免碎片化')

# 3.6
doc.add_heading('3.6 vectordb/store.py — ChromaDB 向量存储', level=2)
add_code('''class VectorStore:
    """Manages document embeddings with ChromaDB."""

    def __init__(self, collection_name="knowledge_base"):
        self.client = chromadb.PersistentClient(
            path=str(config.vectordb_dir),           # SQLite 持久化
            settings=ChromaSettings(anonymized_telemetry=False),
        )
        self.collection_name = collection_name
        self._ensure_collection()

    def _ensure_collection(self):
        """获取或创建 collection"""
        try:
            self.collection = self.client.get_collection(self.collection_name)
        except Exception:
            self.collection = self.client.create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"},  # 余弦距离
            )

    def add(self, doc_id, chunks, embeddings, metadatas):
        """批量添加 chunk，返回 chunk ID 列表"""
        chunk_ids = [f"{doc_id}_chunk_{i}" for i in range(len(chunks))]
        self.collection.add(
            ids=chunk_ids,
            documents=chunks,
            embeddings=embeddings,
            metadatas=metadatas,
        )
        return chunk_ids

    def search(self, query_embedding, top_k=None, where=None):
        """语义搜索：输入向量，返回最相似的 chunk"""
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k or config.dense_top_k,
            where=where,  # 支持元数据过滤（如限定文档范围）
            include=["documents", "metadatas", "distances"],
        )
        hits = []
        if results["ids"] and results["ids"][0]:
            for i, chunk_id in enumerate(results["ids"][0]):
                hits.append({
                    "chunk_id": chunk_id,
                    "doc_id": self._extract_doc_id(chunk_id),
                    "text": results["documents"][0][i],
                    "metadata": results["metadatas"][0][i],
                    # ChromaDB 返回 cosine distance，转换回 similarity
                    "score": 1.0 - results["distances"][0][i],
                })
        return hits''')

add_paragraph('关键设计：')
add_bullet('PersistentClient：数据自动持久化到 SQLite，无需手动 save')
add_bullet('hnsw:space: cosine：使用 HNSW 近似最近邻算法，O(log N) 时间复杂度')
add_bullet('score = 1.0 - distance：将余弦距离转换为相似度分数（越高越好）')
add_bullet('where 参数：支持按 doc_id 等元数据过滤，实现限定文档范围的检索')
add_bullet('chunk_id 命名规则：{doc_id}_chunk_{index}，方便反查所属文档')

# 3.7
doc.add_heading('3.7 retriever/hybrid.py — 混合检索 + BM25 + RRF', level=2)

add_paragraph('BM25 实现 — 核心公式：', bold=True)
add_code('''score(d, q) = Σ IDF(qi) × TF(qi, d) × (k1+1) / (TF(qi, d) + k1 × (1-b + b×|d|/avgdl))

参数说明：
  k1 = 1.5   控制词频饱和度（TF 的影响上限）
  b  = 0.75  控制文档长度归一化（惩罚长文档）
  avgdl      所有文档的平均长度''')

add_code('''class BM25Scorer:
    """Minimal BM25 implementation for sparse retrieval."""

    def __init__(self, k1=1.5, b=0.75):
        self.k1 = k1
        self.b = b
        self.corpus: List[str] = []            # 所有文档文本
        self.doc_freq: Dict[str, int] = {}      # DF: 每个词出现在多少文档中
        self.avgdl: float = 0                   # 平均文档长度
        self.N: int = 0                         # 文档总数

    def index(self, chunks):
        """构建 BM25 索引"""
        self.corpus = [c["text"] for c in chunks]
        self._chunk_meta = chunks
        total_length = 0
        for text in self.corpus:
            tokens = self._tokenize(text)
            total_length += len(tokens)
            # DF: 每篇文档对一个词只计一次
            for token in set(tokens):
                self.doc_freq[token] = self.doc_freq.get(token, 0) + 1
        self.avgdl = total_length / max(len(self.corpus), 1)
        self.N = len(self.corpus)

    def _score(self, query_tokens, doc_text, doc_idx):
        """BM25 打分"""
        doc_tokens = self._tokenize(doc_text)
        doc_len = len(doc_tokens)
        if doc_len == 0:
            return 0.0
        # 计算文档内每个词的 TF
        tf = {}
        for t in doc_tokens:
            tf[t] = tf.get(t, 0) + 1

        score = 0.0
        for token in query_tokens:
            if token not in self.doc_freq:
                continue
            df = self.doc_freq[token]
            # IDF: 逆文档频率（平滑版）
            idf = np.log(1.0 + (self.N - df + 0.5) / (df + 0.5))
            t = tf.get(token, 0)
            # BM25 核心公式
            numerator = t * (self.k1 + 1)
            denominator = t + self.k1 * (1 - self.b + self.b * doc_len / self.avgdl)
            score += idf * numerator / denominator
        return score''')

add_paragraph('CJK 感知分词：', bold=True)
add_code('''@staticmethod
def _tokenize(text: str) -> List[str]:
    """CJK-aware tokenization"""
    text = text.lower()
    tokens = []
    for char in text:
        # 保留：字母、数字、CJK统一表意文字、日文假名
        if char.isalnum() or '一' <= char <= '鿿' or '぀' <= char <= 'ヿ':
            tokens.append(char)
        elif char.isspace():
            tokens.append(char)
    joined = "".join(tokens)
    # 按空格分割，过滤单字符（减少噪音）
    return [t for t in re.split(r'\\s+', joined) if t and len(t) > 1]''')

add_paragraph('RRF（Reciprocal Rank Fusion）融合算法：', bold=True)
add_code('''def _rrf_fusion(self, dense, sparse, k=60):
    """RRF 融合：合并 Dense 和 Sparse 的排序结果"""
    scores: Dict[str, float] = {}
    docs: Dict[str, dict] = {}

    # Dense 检索结果
    for rank, hit in enumerate(dense):
        cid = hit["chunk_id"]
        scores[cid] = scores.get(cid, 0) + 1.0 / (k + rank + 1)
        docs[cid] = hit
        docs[cid]["source"] = "dense"

    # Sparse 检索结果
    for rank, hit in enumerate(sparse):
        cid = hit["chunk_id"]
        scores[cid] = scores.get(cid, 0) + 1.0 / (k + rank + 1)
        if cid not in docs:
            docs[cid] = hit
            docs[cid]["source"] = "sparse"
        else:
            docs[cid]["source"] = "both"  # 两个检索器都命中

    # 按 RRF 分数降序排列
    sorted_ids = sorted(scores, key=scores.get, reverse=True)
    return [{**docs[cid], "rrf_score": round(scores[cid], 6)} for cid in sorted_ids]''')

add_paragraph('RRF 公式: score(chunk) = Σ 1/(k + rank_i)，k=60（平滑参数）。source 字段含义：dense（仅向量命中）、sparse（仅BM25命中）、both（双命中，最可靠）。')

# 3.8
doc.add_heading('3.8 llm.py — 多后端 LLM 抽象', level=2)
add_paragraph('策略模式 + 适配器模式实现多 LLM 后端的统一接口。')

add_paragraph('抽象基类：', bold=True)
add_code('''class LLMBackend:
    """Abstract base for LLM backends."""
    def generate(self, system, messages, max_tokens, temperature) -> str:
        raise NotImplementedError
    async def generate_stream(self, system, messages, max_tokens, temperature):
        raise NotImplementedError
    @property
    def name(self) -> str:
        raise NotImplementedError''')

add_paragraph('Anthropic 后端（原生支持 system prompt 独立参数）：', bold=True)
add_code('''class AnthropicBackend(LLMBackend):
    def generate(self, system, messages, max_tokens=2048, temperature=0.3):
        from anthropic import Anthropic
        client = Anthropic(api_key=self.api_key)
        resp = client.messages.create(
            model=self.model,
            max_tokens=max_tokens,
            system=system,           # Anthropic 原生独立 system 参数
            messages=messages,
            temperature=temperature,
        )
        return resp.content[0].text

    async def generate_stream(self, system, messages, max_tokens=2048, temperature=0.3):
        from anthropic import AsyncAnthropic
        client = AsyncAnthropic(api_key=self.api_key)
        async with client.messages.stream(
            model=self.model, system=system, messages=messages,
            max_tokens=max_tokens, temperature=temperature,
        ) as stream:
            async for event in stream:
                if event.type == "content_block_delta":
                    yield event.delta.text
                elif event.type == "message_stop":
                    break''')

add_paragraph('OpenAI 兼容后端（统一处理 OpenAI / DeepSeek / Ollama）：', bold=True)
add_code('''class OpenAICompatibleBackend(LLMBackend):
    """OpenAI / DeepSeek / Ollama — 均使用 OpenAI 兼容 API"""

    def _build_payload(self, system, messages, max_tokens, temperature):
        # 关键差异: OpenAI 格式不区分 system/messages，需合并
        full_messages = [{"role": "system", "content": system}] + list(messages)
        return {
            "model": self.model,
            "messages": full_messages,
            "max_tokens": max_tokens,
            "temperature": temperature,
        }

    def generate(self, system, messages, max_tokens=2048, temperature=0.3):
        from openai import OpenAI
        client = OpenAI(api_key=self.api_key, base_url=self.base_url)
        payload = self._build_payload(system, messages, max_tokens, temperature)
        payload["stream"] = False
        resp = client.chat.completions.create(**payload)
        return resp.choices[0].message.content or ""

    async def generate_stream(self, system, messages, max_tokens=2048, temperature=0.3):
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=self.api_key, base_url=self.base_url)
        payload = self._build_payload(system, messages, max_tokens, temperature)
        payload["stream"] = True
        stream = await client.chat.completions.create(**payload)
        async for chunk in stream:
            delta = chunk.choices[0].delta if chunk.choices else None
            if delta and delta.content:
                yield delta.content''')

add_paragraph('自动检测逻辑（责任链模式）：', bold=True)
add_code('''class LLM:
    """Unified LLM interface with auto-detection."""

    def _resolve(self, backend, model, api_key, base_url):
        """Auto-detect the best available backend."""
        bt = (backend or os.environ.get("LLM_BACKEND", "")).lower()

        # ── 1. 显式指定 LLM_BACKEND 环境变量 ──
        if bt == "ollama":
            ollama_model = model or self._detect_ollama_model()
            return OpenAICompatibleBackend(model=ollama_model, ...)
        if bt == "anthropic" or bt == "claude":
            return AnthropicBackend(...)
        if bt == "deepseek":
            return OpenAICompatibleBackend(model="deepseek-chat", ...)
        if bt == "openai":
            return OpenAICompatibleBackend(model="gpt-4o-mini", ...)

        # ── 2. 自动检测 Ollama（本地可达，免费优先）──
        if self._ollama_available():
            return OpenAICompatibleBackend(model=best_ollama_model, ...)

        # ── 3. 检测 Anthropic API Key ──
        if os.environ.get("ANTHROPIC_API_KEY"):
            return AnthropicBackend(model="claude-sonnet-4-6", ...)

        # ── 4. 检测 DeepSeek API Key ──
        if os.environ.get("DEEPSEEK_API_KEY"):
            return OpenAICompatibleBackend(model="deepseek-chat", ...)

        # ── 5. 检测 OpenAI API Key ──
        if os.environ.get("OPENAI_API_KEY"):
            return OpenAICompatibleBackend(model="gpt-4o-mini", ...)

        # ── 6. 全部未配置 → 抛出错误 ──
        raise RuntimeError("No LLM backend available. ...")

    @staticmethod
    def _ollama_available() -> bool:
        """通过 HTTP GET localhost:11434/api/tags 检测 Ollama 是否运行"""
        try:
            url = (os.environ.get("OLLAMA_HOST", "http://localhost:11434")).rstrip("/v1") + "/api/tags"
            with urllib.request.urlopen(url, timeout=2) as resp:
                data = json.loads(resp.read())
            return len(data.get("models", [])) > 0
        except Exception:
            return False''')

add_paragraph('设计模式总结：策略模式（不同后端实现同一接口）+ 适配器模式（OpenAI兼容后端统一多个API）+ 工厂方法（_resolve 根据环境创建正确后端）。自动检测优先级：显式配置 > Ollama > Anthropic > DeepSeek > OpenAI。')

# 3.9
doc.add_heading('3.9 agent/agent.py — 知识问答 Agent', level=2)
add_paragraph('核心编排层：检索 + 上下文组装 + LLM 调用 + 会话管理。')

add_code('''class KnowledgeAgent:
    """Orchestrates retrieval + LLM generation."""

    def __init__(self, retriever, llm):
        self.retriever = retriever
        self.llm = llm
        self._conversations: Dict[str, List[dict]] = {}  # 会话字典

    def query(self, question, conv_id="default", top_k=None, doc_ids=None):
        """同步问答（等待完整回复后返回）"""
        # 1. 检索
        chunks = self.retriever.retrieve(question, top_k=top_k, doc_ids=doc_ids)
        if not chunks:
            return {"answer": "No relevant documents found.", "sources": [], ...}

        # 2. 构建上下文
        context = self._build_context(chunks)
        # 3. 填充 System Prompt
        system_prompt = SYSTEM_PROMPT.format(context=context)
        # 4. 获取会话历史 + 拼接新消息
        messages = self._get_history(conv_id) + [{"role": "user", "content": question}]
        # 5. 调用 LLM
        answer = self.llm.generate(system=system_prompt, messages=messages,
                                    max_tokens=config.llm_max_tokens, temperature=0.3)
        # 6. 更新会话历史
        self._add_to_history(conv_id, "user", question)
        self._add_to_history(conv_id, "assistant", answer)
        return {"answer": answer, "sources": self._extract_sources(chunks), ...}

    async def query_stream(self, question, conv_id="default", ...):
        """流式问答（SSE）"""
        chunks = self.retriever.retrieve(question, ...)

        # 先发送 sources
        sources = self._extract_sources(chunks)
        yield json.dumps({"type": "sources", "data": sources})

        # 流式发送 token
        full_answer = ""
        async for token in self.llm.generate_stream(...):
            full_answer += token
            yield json.dumps({"type": "token", "data": token})

        # 保存历史 + 发送结束
        self._add_to_history(conv_id, "user", question)
        self._add_to_history(conv_id, "assistant", full_answer)
        yield json.dumps({"type": "done", "data": {"conv_id": conv_id, ...}})

    def _build_context(self, chunks):
        """格式化检索结果为 LLM 可读的上下文"""
        parts = []
        for i, chunk in enumerate(chunks):
            filename = chunk["metadata"]["filename"]
            score = chunk.get("rrf_score") or chunk.get("score", 0)
            parts.append(
                f"--- Chunk {i+1} [source: {filename}, index: {cidx}, "
                f"relevance: {score:.4f}] ---\\n{chunk['text']}\\n"
            )
        return "\\n".join(parts)''')

add_paragraph('会话管理：')
add_bullet('内存字典 conv_id → [{role, content}, ...]')
add_bullet('LRU 策略：超过 100 个会话时删除字典中最旧的 key')
add_bullet('支持手动 reset：_conversations.pop(conv_id, None)')

add_paragraph('来源提取：')
add_code('''@staticmethod
def _extract_sources(chunks):
    """按文档分组汇聚来源信息"""
    seen = {}
    for c in chunks:
        filename = c["metadata"]["filename"]
        doc_id = c["doc_id"]
        if doc_id not in seen:
            seen[doc_id] = {"doc_id": doc_id, "filename": filename, "chunks": []}
        seen[doc_id]["chunks"].append({
            "chunk_id": c["chunk_id"],
            "chunk_index": c["metadata"]["chunk_index"],
            "score": c.get("rrf_score") or c.get("score", 0),
            "text_preview": c["text"][:200] + "...",
        })
    return list(seen.values())''')

# 3.10
doc.add_heading('3.10 agent/prompts.py — 提示词设计', level=2)
add_code('''SYSTEM_PROMPT = """You are an AI Knowledge Base Agent.
You answer user questions based on the provided document context.

## Rules
1. Answer ONLY based on the provided context chunks. If the context does
   not contain enough information, say "I don't have enough information
   in the knowledge base to answer this question."
2. ALWAYS cite sources by referencing the chunk metadata (filename, chunk
   index) when providing information.
3. Format citations like [source: filename, chunk N].
4. Be concise and precise. Use bullet points for lists.
5. If the user asks a question in Chinese, respond in Chinese. Otherwise
   use the same language as the user.
6. When quoting directly from a document, use quotation marks and cite
   the source.
7. If the context is fragmented, synthesize the information into a
   coherent answer.
8. NEVER make up information that is not in the context.

## Current Context
The following chunks were retrieved from the knowledge base:

{context}

## Conversation Guidelines
- Answer the user's question using the context above
- Cite sources for every factual claim
- If the answer spans multiple chunks, synthesize them naturally
"""''')
add_paragraph('设计要点：')
add_bullet('规则1和8是反幻觉的核心约束：只看上下文，不编造信息')
add_bullet('规则2和3强制可溯源：每条事实引用来源文件和chunk编号')
add_bullet('规则5是中英文自适应：用户用中文提问就用中文回答')
add_bullet('{context} 占位符在运行时由 agent.py 的 _build_context() 动态填充')

# 3.11
doc.add_heading('3.11 kg/builder.py — 知识图谱', level=2)
add_paragraph('两种实体抽取模式：')

add_code('''class KnowledgeGraphBuilder:
    """Build a knowledge graph from document chunks."""

    def build(self, use_llm=False):
        if use_llm and config.anthropic_api_key:
            self._build_with_llm(chunks)       # LLM 精确抽取（慢但有费用）
        else:
            self._build_with_keywords(chunks)  # 关键词规则抽取（快且免费）

    def _build_with_keywords(self, chunks):
        """基于规则的关键词抽取"""
        for chunk in chunks:
            entities = self._extract_keywords(chunk["text"])
            # 统计词频，取 top-8
            ...
        # 构建二部图: 文档节点 ←→ 实体节点

    @staticmethod
    def _extract_keywords(text):
        keywords = []
        # 英文: 提取大写短语 (如 "Machine Learning")
        english = re.findall(r'\\b[A-Z][a-z]+(?:\\s+[A-Z][a-z]+)*\\b', text)
        keywords.extend([p.lower() for p in english if len(p) > 5])
        # CJK: 提取 2-4 字序列，过滤停用词
        cjk = re.findall(r'[一-鿿]{2,4}', text)
        stop = {'这是', '一个', '可以', '我们', '这个', '那个', ...}
        keywords.extend([c for c in cjk if c not in stop])
        return keywords

    def _build_with_llm(self, chunks):
        """用 Claude 提取实体（前50个chunk）"""
        for chunk in chunks[:50]:
            entities = self._llm_extract_entities(client, chunk["text"])
            # 实体类型: person/organization/location/product/technology/concept

    def to_cytoscape(self):
        """导出为 Cytoscape.js 兼容的 JSON 格式"""
        elements = []
        for node_id, data in self.graph.nodes(data=True):
            elements.append({"data": {"id": node_id, "label": ..., "type": ...}})
        for u, v, data in self.graph.edges(data=True):
            elements.append({"data": {"source": u, "target": v, "weight": ...}})
        return {"elements": elements}''')

add_paragraph('两种模式对比：')
add_table(
    ['维度', '关键词模式', 'LLM 模式'],
    [
        ['速度', '毫秒级', '秒级（需 API 调用）'],
        ['费用', '免费', '按 token 计费'],
        ['准确性', '一般（基于规则）', '高（语义理解）'],
        ['中文支持', '有限', '好'],
        ['依赖', '无外部依赖', '需要 LLM API'],
    ]
)

# 3.12
doc.add_heading('3.12 main.py — FastAPI 路由', level=2)
add_paragraph('全部 11 个 API 端点汇总：')
add_table(
    ['方法', '路径', '功能'],
    [
        ['GET', '/api/health', '健康检查 + chunk/doc 数量'],
        ['POST', '/api/upload', '上传文档（multipart/form-data）'],
        ['POST', '/api/upload-url', '抓取网页并索引'],
        ['GET', '/api/documents', '列出已索引文档'],
        ['DELETE', '/api/documents/{id}', '删除文档及所有 chunk'],
        ['POST', '/api/chat', '同步问答'],
        ['POST', '/api/chat/stream', '流式问答（SSE）'],
        ['POST', '/api/conversation/{id}/reset', '清除会话历史'],
        ['GET', '/api/search?q=...&top_k=N', '纯检索（不使用 LLM）'],
        ['GET', '/api/knowledge-graph', '知识图谱 Cytoscape JSON'],
        ['GET', '/api/knowledge-graph/stats', '图谱统计'],
    ]
)

add_paragraph('模块级初始化（应用启动时执行一次）：')
add_code('''# ── 初始化（模块级别，import 时执行）──
embedder = Embedder()                          # 嵌入模型（懒加载）
vector_store = VectorStore()                   # ChromaDB 连接
retriever = HybridRetriever(vector_store, embedder)  # 混合检索器
retriever.build_bm25_index()                   # 从已有数据重建 BM25 索引

try:
    llm = LLM()                                # 自动检测 LLM 后端
except RuntimeError:
    llm = None                                 # 无 LLM 时 chat 不可用，但检索仍可用

agent = KnowledgeAgent(retriever, llm) if llm else None
chunker = TextChunker()
kg_builder = KnowledgeGraphBuilder(vector_store)

app = FastAPI(title="AI Knowledge Base Agent", version="1.0.0")''')

add_paragraph('文档上传的完整索引流程：')
add_code('''def _index_document(doc) -> dict:
    """通用索引流水线"""
    # 1. 重复检测：按 source 字段判断
    for ex in vector_store.list_documents():
        if ex.get("source") == doc.metadata.get("source"):
            return {"status": "already_indexed", ...}

    # 2. 文本切分
    chunks = chunker.split(doc.content)

    # 3. 批量向量化
    embeddings = embedder.embed_batch(chunks)

    # 4. 构建元数据（每个 chunk 带有 doc_id, filename, chunk_index 等）
    metadatas = [{...} for i, chunk in enumerate(chunks)]

    # 5. 存储到 ChromaDB
    vector_store.add(doc.id, chunks, embeddings, metadatas)

    # 6. 全量重建 BM25 索引
    retriever.build_bm25_index()

    return {"status": "indexed", "doc_id": doc.id, "chunks": len(chunks), ...}''')

# 3.13
doc.add_heading('3.13 前端 app.js — SSE 流式处理', level=2)
add_paragraph('SSE 流式问答的核心实现：')
add_code('''async function sendMessage() {
    const question = input.value.trim();
    // 显示用户消息
    addMessage('user', question);
    // 创建 assistant 占位（loading 动画）
    const assistantMsg = addMessage('assistant', '', true);

    const resp = await fetch(`${API}/chat/stream`, {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ question, conv_id: convId }),
    });

    const reader = resp.body.getReader();      // ReadableStream API
    const decoder = new TextDecoder();
    let buffer = '';

    while (true) {
        const { done, value } = await reader.read();
        if (done) break;

        buffer += decoder.decode(value, { stream: true });
        const lines = buffer.split('\\n');
        buffer = lines.pop();  // 保留不完整行（处理跨 chunk 的 SSE 消息）

        for (const line of lines) {
            if (!line.startsWith('data: ')) continue;
            const data = JSON.parse(line.slice(6));  // 去掉 "data: " 前缀
            // data.type: "sources" | "token" | "done"
            if (data.type === 'token') {
                contentEl.textContent += data.data;   // 逐 token 追加到界面
                scrollToBottom();
            } else if (data.type === 'sources') {
                sourceData = data.data;
            }
        }
    }
    // 流结束后展示来源面板
    if (sourceData) showSources(sourceData);
}''')

add_paragraph('buffer 机制详解：SSE 协议以双换行（\\n\\n）为消息分隔符。当 ReadableStream 的 chunk 边界恰好落在一条 SSE 消息中间时，lines.pop() 将不完整的最后一行保存到 buffer，等待下一次 read() 拼接。')

add_paragraph('知识图谱渲染：')
add_code('''function renderGraph(data) {
    cyInstance = cytoscape({
        container: document.getElementById('cyContainer'),
        elements: data.elements,   // 来自 /api/knowledge-graph 的 Cytoscape JSON
        style: [
            { selector: 'node[type="document"]',
              style: { 'background-color': '#10b981', 'font-weight': 'bold' } },
            { selector: 'node[type="entity"]',
              style: { 'background-color': '#4f46e5' } },
            { selector: 'edge',
              style: { 'width': 'mapData(weight, 1, 10, 0.5, 3)' } },
        ],
        layout: {
            name: 'cose',              // 力导向布局
            nodeRepulsion: 4000,       // 节点排斥力
            idealEdgeLength: 80,       // 理想边长
        },
    });
}''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  PART 4: INTERVIEW Q&A (30 questions)
# ═══════════════════════════════════════════════════════════
doc.add_heading('四、面试问题与答案（30 题）', level=1)

# ── 4.1 RAG 基础 ──
doc.add_heading('4.1 RAG 基础', level=2)

doc.add_heading('Q1: 什么是 RAG？为什么需要 RAG？', level=3)
add_paragraph('RAG（Retrieval-Augmented Generation，检索增强生成）是一种结合检索和生成的 AI 架构。核心流程：用户提问 → 从知识库检索相关文档片段 → 将检索结果作为上下文注入 LLM → LLM 基于上下文生成答案。')
add_paragraph('为什么需要 RAG？')
add_bullet('LLM 有知识截止日期，无法回答最新信息（如 2024 年发生的新闻）')
add_bullet('LLM 会产生幻觉（Hallucination），编造不存在的事实')
add_bullet('企业有大量私有文档（内部报告、手册等），LLM 训练数据中没有')
add_bullet('RAG 提供可溯源（Grounding）的答案：每条事实可以追溯到原文具体位置')

doc.add_heading('Q2: 项目中文本切分的 chunk_size 为什么选 500？', level=3)
add_paragraph('这是嵌入质量与检索精度之间的平衡：')
add_bullet('太小（<200 字符）：语义碎片化，一个完整句子可能被切成多个 chunk。检索时单个 chunk 的上下文不足以支撑 LLM 理解，embedding 质量也会下降')
add_bullet('太大（>1000 字符）：一个 chunk 包含过多无关信息，检索精度下降。即使检索到相关 chunk，也会浪费 LLM 的上下文窗口在无关内容上')
add_bullet('500 是一个经验平衡点：约等于 2-3 个中文段落，既能保持语义完整性，又能在检索时精确定位相关信息')
add_bullet('实际中应根据文档类型调优：代码文档使用更小的 chunk（~300 字符），学术论文使用更大的 chunk（~800 字符）')

doc.add_heading('Q3: chunk_overlap 的作用是什么？为什么是 50？', level=3)
add_paragraph('chunk_overlap 让相邻 chunk 共享一部分文本（尾部和头部重叠），防止关键信息刚好落在 chunk 边界而被切断。')
add_paragraph('举例："苹果公司于 2024 年发布了 iPhone 16，搭载了 A18 芯片。"')
add_bullet('无 overlap：Chunk 1 "苹果公司于 2024 年发布了" | Chunk 2 "iPhone 16，搭载了 A18 芯片。"——"发布了" 和 "iPhone 16" 被切断')
add_bullet('有 overlap（50字）：Chunk 1 "...2024 年发布了 iPhone 16" | Chunk 2 "发布了 iPhone 16，搭载了 A18 芯片。"——"发布了 iPhone 16" 在两边都存在，LLM 能完整理解')
add_paragraph('50 是 chunk_size（500）的 10%，经实验验证的合理值：提供足够的上下文连续性，又不会导致太多冗余存储。')

# ── 4.2 嵌入与向量检索 ──
doc.add_heading('4.2 嵌入与向量检索', level=2)

doc.add_heading('Q4: 为什么选择 all-MiniLM-L6-v2 而不是其他模型？', level=3)
add_table(
    ['因素', 'all-MiniLM-L6-v2', 'OpenAI text-embedding-3', 'BGE-M3（多语言）'],
    [
        ['向量维度', '384', '1536 / 512', '1024'],
        ['模型大小', '~80MB', '需 API 调用', '~2GB'],
        ['单条延迟', '~5ms（CPU）', '~50ms（网络）', '~15ms（GPU）'],
        ['费用', '免费', '按 token 计费', '免费'],
        ['离线可用', '✅', '❌', '✅'],
        ['多语言', '一般', '好', '优秀'],
    ]
)
add_paragraph('选择理由：本地运行零成本零延迟；384 维在文档 QA 场景与 1536 维差距 < 5%；80MB 首次下载快；CPU 推理毫秒级。')

doc.add_heading('Q5: 向量检索中的余弦相似度是怎么回事？为什么需要归一化？', level=3)
add_paragraph('余弦相似度衡量两个向量在方向上的相似程度：cos(θ) = (A·B) / (|A| × |B|)')
add_paragraph('当向量经过 L2 归一化后（模长为 1），|A| = |B| = 1，余弦相似度退化为简单的点积运算：cos(θ) = A·B。这大幅简化了向量数据库中的相似度计算。')
add_paragraph('在 ChromaDB 中，存储的是余弦距离：distance = 1 - cosine_sim。所以我们在代码中用 score = 1.0 - distance 转换回相似度分数。')

doc.add_heading('Q6: 为什么用 ChromaDB 而不是 FAISS 或 Milvus？', level=3)
add_table(
    ['维度', 'ChromaDB', 'FAISS', 'Milvus'],
    [
        ['部署方式', '嵌入式（pip install）', '嵌入式（C++ 库）', '需要独立服务'],
        ['持久化', '自动（SQLite）', '手动（需自己序列化）', '自动'],
        ['元数据过滤', '原生支持 where 子句', '不支持', '原生支持'],
        ['适用规模', '< 100万 chunk', '> 1000万 chunk', '> 100万 chunk'],
        ['运维复杂度', '零（零配置）', '低', '中高（需维护服务）'],
        ['API 风格', 'Pythonic', 'C++ 风格', 'gRPC'],
    ]
)
add_paragraph('ChromaDB 最适合本项目"个人/小团队知识库"的定位：单机部署，零运维，开箱即用。')

doc.add_heading('Q7: HNSW 是什么？为什么快？', level=3)
add_paragraph('HNSW（Hierarchical Navigable Small World）是近似最近邻搜索（ANN）算法，被 ChromaDB 内置使用。')
add_paragraph('原理：构建多层图结构。顶层节点稀疏，用于长距离"跳跃"快速定位大致区域；底层节点稠密，用于精确搜索。')
add_paragraph('搜索过程：从顶层开始 → 在每层找到最近节点 → 下降到下一层 → 重复直到最底层 → 在底层做精确搜索。')
add_paragraph('时间复杂度 O(log N)，对比暴力搜索 O(N)。代价是牺牲少量精度（通常 < 1%），换取百倍以上的速度提升。在 10,000 条 chunk 中，暴力搜索需要计算 10,000 次距离，HNSW 只需要约 50-200 次。')

# ── 4.3 BM25 与混合检索 ──
doc.add_heading('4.3 BM25 与混合检索', level=2)

doc.add_heading('Q8: 解释 BM25 公式的每个部分。', level=3)
add_paragraph('BM25 完整公式：score(d, q) = Σ IDF(qi) × (k1+1)×TF / (TF + k1×(1-b + b×|d|/avgdl))')
add_table(
    ['参数', '含义', '为什么这样设置'],
    [
        ['IDF(qi)', '逆文档频率', '稀有词权重高（如"量子计算"），常见词权重低（如"的"、"是"）'],
        ['TF', '词频（Term Frequency）', '出现次数越多，该词对文档越重要'],
        ['k1 = 1.5', '词频饱和度参数', '控制 TF 的影响上限。防止一个词出现 100 次就完全主导分数'],
        ['b = 0.75', '长度归一化参数', '惩罚长文档。0.75 表示 75% 的长度因素被归一化，比完全归一化（b=1）保留一些优势给长文档'],
        ['avgdl', '所有文档的平均长度', '用于长度归一化的基准值'],
    ]
)
add_paragraph('直觉理解：查询中稀有的词语权重高；文档中出现多次的词语加分；但出现太多不会线性加分（k1 控制）；长文档不会因为长度而天然占优（b 控制）。')

doc.add_heading('Q9: 为什么 Dense + Sparse 混合检索比单一检索更好？', level=3)
add_paragraph('纯向量检索（Dense）的弱点：')
add_bullet('对专有名词不敏感：产品编号"SKU-A12345"、错误码"ERR_TIMEOUT"在向量空间中可能与不相关文本距离很近')
add_bullet('"Apple 公司" 和 "apple 水果" 在向量空间中距离很近，但语义完全不同')
add_paragraph('纯 BM25（Sparse）的弱点：')
add_bullet('只能做字面匹配，不理解语义："汽车" 和 "轿车" 在 BM25 中是两个完全不同的 token')
add_bullet('无法处理同义词替换："如何提升性能" vs "怎么让系统更快"——没有共同词汇，BM25 得分为 0')
add_paragraph('混合检索的互补效果：BM25 负责精确关键词匹配（产品编号、API 名称等）；Dense 负责语义理解（近义词、改写后的查询等）；两者结合后，无论用户怎么表达都能找到相关内容。')

doc.add_heading('Q10: RRF 相比加权求和的优势是什么？为什么 k=60？', level=3)
add_paragraph('加权求和的问题：需要调 α 参数（α×dense_score + (1-α)×sparse_score），两个 score 的量纲完全不同（向量相似度 0-1，BM25 分数无上限），需要额外做 score 归一化。α=0.3 还是 0.7？不同查询场景可能需要不同的 α。')
add_paragraph('RRF 的优势：不关心原始分数的量纲，只关心相对排名；无需任何超参数调优；对排名靠前的结果更敏感（1/(k+1) >> 1/(k+20)）；两个检索器结果集有重叠时效果最好，已被多篇学术论文验证 MAP 提升 5-15%。')
add_paragraph('k=60 来自学术界（Cornell 大学、Elasticsearch 团队的论文验证），其作用是平滑排名差异——k 越大，排名靠后的结果权重越接近排名靠前的。60 是一个"使排名靠前但并非第一的结果也得到合理权重"的平衡值。')

# ── 4.4 LLM 相关 ──
doc.add_heading('4.4 LLM 相关', level=2)

doc.add_heading('Q11: 项目中 LLM 自动检测是怎么实现的？', level=3)
add_paragraph('使用责任链模式（Chain of Responsibility），从免费到付费逐级检测：')
add_bullet('1. 检查环境变量 LLM_BACKEND 是否显式指定（如 "ollama" / "anthropic"）')
add_bullet('2. HTTP GET localhost:11434/api/tags 检测 Ollama 是否运行且至少有 1 个模型 → 选择最优模型（qwen > llama > deepseek > 第一个可用）')
add_bullet('3. 检查 ANTHROPIC_API_KEY 是否设置 → 使用 Claude')
add_bullet('4. 检查 DEEPSEEK_API_KEY 是否设置 → 使用 DeepSeek')
add_bullet('5. 检查 OPENAI_API_KEY 是否设置 → 使用 OpenAI')
add_bullet('6. 全部未配置 → 抛出 RuntimeError，chat 功能禁用但检索仍可用')

doc.add_heading('Q12: 为什么把 system prompt 放在检索上下文里而不是 fine-tune 模型？', level=3)
add_paragraph('成本：Fine-tune 需要大量高质量标注数据（至少几百条）和 GPU 算力，远高于 prompt engineering 的成本。灵活性：修改 prompt 即时生效（调整 citations 格式、改变回答风格、添加新规则），fine-tune 需要重新训练。模型无关：prompt 可以在任何 LLM（Claude/GPT/开源模型）上使用，fine-tune 绑定了特定模型。可解释性：prompt 是透明的，可以看到给 LLM 的具体指令；fine-tune 是黑盒的。')

doc.add_heading('Q13: 如何防止 LLM 幻觉（Hallucination）？', level=3)
add_paragraph('项目中采用了多层防护策略：')
add_bullet('1. System Prompt 硬约束："Answer ONLY based on the provided context." "NEVER make up information."')
add_bullet('2. 明确的上下文不足策略：当检索结果不包含相关答案时，LLM 被告知要说 "I don\'t have enough information"')
add_bullet('3. 强制引用机制：每条事实性陈述必须使用 [source: filename, chunk N] 格式标注出处')
add_bullet('4. 低温度设置：temperature=0.3 降低输出的随机性，减少"创造性编造"')
add_bullet('5. 如果需要更进一步，可以加入 RAGAS 评估框架，量化检测幻觉率（Faithfulness 指标）')

doc.add_heading('Q14: 流式回答（SSE）怎么实现的？有什么注意事项？', level=3)
add_paragraph('服务端：FastAPI 通过 StreamingResponse + async generator 实现。每次 yield "data: {json}\\n\\n"，SSE 协议以双换行为消息分隔符。')
add_paragraph('前端：使用浏览器原生 ReadableStream API（resp.body.getReader()）逐块读取，通过 buffer 机制处理跨 chunk 的不完整消息。')
add_paragraph('注意事项：SSE 是单向通信（服务器→客户端），客户端无法中途打断已发出的消息；生产环境需配置反向代理（Nginx）关闭缓冲（proxy_buffering off），否则 token 会被缓冲导致非实时推送；需要考虑断线重连机制（SSE 原生支持 Last-Event-ID）。')

doc.add_heading('Q15: 为什么要做多后端 LLM 抽象？用了什么设计模式？', level=3)
add_paragraph('原因：不同场景需要不同的 LLM 后端——开发测试用免费的 Ollama，生产环境用 Claude（质量最高），中文场景用 DeepSeek（性价比高）。用户可以根据自己的需求和预算灵活切换。')
add_paragraph('设计模式：')
add_bullet('策略模式（Strategy）：LLMBackend 抽象基类定义统一接口（generate / generate_stream），AnthropicBackend 和 OpenAICompatibleBackend 各自实现')
add_bullet('适配器模式（Adapter）：OpenAICompatibleBackend 将 OpenAI / DeepSeek / Ollama 三个不同 API 统一为一个接口。关键处理：Anthropic API 原生支持 system 参数独立传入，OpenAI 格式需要将 system 合并到 messages[0]')
add_bullet('工厂方法（Factory Method）：LLM._resolve() 根据环境变量和检测结果自动创建正确的后端实例。上层 KnowledgeAgent 完全不感知使用的是哪个 LLM')

# ── 4.5 知识图谱 ──
doc.add_heading('4.5 知识图谱', level=2)

doc.add_heading('Q16: 知识图谱在这个项目中的作用是什么？', level=3)
add_bullet('1. 文档关系可视化：展示哪些文档共享相同的关键实体，帮助用户发现跨文档的知识关联')
add_bullet('2. 实体导航：快速了解知识库覆盖了哪些主题/人物/产品/技术概念，便于知识库概览')
add_bullet('3. 二部图结构：文档节点（绿色）↔ 实体节点（紫色），节点大小反映连接度（degree），边粗细反映关联强度')
add_bullet('4. 前端使用 Cytoscape.js 的 cose 力导向布局，支持拖拽、缩放等交互')

doc.add_heading('Q17: 为什么提供两种实体抽取模式（关键词 vs LLM）？', level=3)
add_paragraph('两种模式适用于不同的使用场景：')
add_bullet('关键词模式（默认）：毫秒级响应，完全免费。适合快速浏览知识库全貌，不需要任何 API 调用')
add_bullet('LLM 模式：需要 API 调用，但准确度高。Claude 可以识别 person / organization / location / product / technology / concept 等分类实体，语义理解比规则强得多')
add_bullet('关键词模式的技术细节：英文提取大写短语（如 "Machine Learning"），中文提取 2-4 字 n-gram 并过滤高频停用词')

# ── 4.6 工程实践 ──
doc.add_heading('4.6 工程实践', level=2)

doc.add_heading('Q18: 为什么把 embedding 模型设计成懒加载？', level=3)
add_paragraph('懒加载（Lazy Initialization）的设计意图：')
add_bullet('启动速度：backend 模块被 import 时不会立即加载 80MB 模型文件，FastAPI 服务可以快速启动并响应 /api/health')
add_bullet('内存优化：如果用户只调用 /api/search 和 /api/documents 等非 LLM 接口，Embedder 在首次使用时才加载')
add_bullet('错误延迟：模型下载失败在首次调用时才暴露，不影响服务的基础健康检查')

doc.add_heading('Q19: 为什么要全量重建 BM25 索引？有什么改进方案？', level=3)
add_paragraph('当前实现：每次上传文档后调用 retriever.build_bm25_index()，从 ChromaDB 全量导出所有 chunk 重建 BM25 索引。这样做简单可靠，但文档多时重建耗时 O(N) 增长。')
add_paragraph('改进方案（增量索引）：新增文档时只更新新增 token 的 DF 计数；新文档的 corpus 追加到列表末尾；维护 total_length 和文档数 N，重新计算 avgdl。具体代码大约需要 20 行改写。')

doc.add_heading('Q20: 会话管理为什么要限制 100 个会话？', level=3)
add_paragraph('会话存储在内存字典中（_conversations: Dict[str, List[dict]]），没有外部持久化（如 Redis）。如果不加限制，大量用户长时间对话会导致内存溢出。100 个会话 × 每会话约 20 条消息 × 每条约 500 字符 ≈ 约 1MB 内存，完全可控。')
add_paragraph('采用简单的 LRU 近似策略：字典 key 数量超过 100 时，用 min(self._conversations.keys()) 删除"字典序最小"的 key。更严谨的做法是用 collections.OrderedDict 实现真正的 LRU（按插入顺序），或给每个会话加 last_access 时间戳。')

doc.add_heading('Q21: _clean_text 函数做了什么？为什么重要？', level=3)
add_paragraph('数据清洗是 RAG 系统的关键环节，"垃圾进垃圾出"（Garbage In, Garbage Out）。')
add_bullet('合并多余空行（3+ → 2）：防止 Chunker 在无意义的空白处切分，影响 chunk 语义完整性')
add_bullet('合并多余空格/制表符（3+ → 2）：减少 embedding 的噪声，避免向量被无关空白主导')
add_bullet('移除空字符（\\x00）：来自部分 PDF 的编码问题，空字符会导致 embedding 向量异常')
add_bullet('首尾去空白：减少无效 token，提供干净的输入给下游处理')

doc.add_heading('Q22: DOCX 为什么要做 XML fallback？', level=3)
add_paragraph('python-docx 库对某些特殊格式的 .docx 文件支持不完善：WPS Office 创建的文档使用了不同的 XML 结构；较老版本的 Microsoft Word 格式；包含文本框（Text Box）的文档——python-docx 的 doc.paragraphs 在这些情况下可能返回空。')
add_paragraph('XML fallback 方案利用了 .docx 本质是 ZIP 文件的特性：打开 ZIP → 找到 document*.xml → 用正则提取所有 <w:p> 段落内的 <w:t> 文本。这是从 OOXML 底层格式直接提取，兼容性最好，确保不丢失内容。')

# ── 4.7 系统设计 ──
doc.add_heading('4.7 系统设计', level=2)

doc.add_heading('Q23: 如果要支持 100 万篇文档，这个系统需要做哪些改造？', level=3)
add_table(
    ['组件', '当前方案', '升级方案', '原因'],
    [
        ['向量数据库', 'ChromaDB（嵌入式）', 'Milvus / Qdrant（分布式）', '支持十亿级向量，分布式索引'],
        ['嵌入推理', 'CPU 推理', 'GPU 推理 或 Embedding API', 'CPU 处理 100万×500chunk 不可行'],
        ['BM25 索引', '内存字典', 'Elasticsearch / Lucene', '磁盘索引 + 成熟的分词和检索能力'],
        ['文档处理', '同步上传即处理', '异步任务队列（Celery + Redis）', '大文档不阻塞 HTTP 请求'],
        ['缓存', '无', 'Redis 缓存热点查询', '减少重复检索 + 降低 LLM 费用'],
        ['多租户', '无', '按用户/项目分 Collection', '数据隔离 + 权限管理'],
        ['Reranker', '无', 'Cross-encoder 精排', '在 RRF 后对 top-20 重排序，提升精度'],
        ['监控', '无', 'Prometheus + Grafana', '检索延迟、召回率、LLM 调用统计'],
    ]
)

doc.add_heading('Q24: 这个项目的核心性能瓶颈在哪里？', level=3)
add_bullet('1. Embedding 生成：CPU 推理每条 chunk 约 5ms。上传包含 200 个 chunk 的文档需要约 1 秒纯嵌入时间。1000 chunks = 5 秒。解决方案：GPU 推理 / 批量 API / 异步处理')
add_bullet('2. BM25 全量重建：每次上传都 O(N) 重新遍历所有 chunk。1万 chunks 时重建约 50ms，可接受；百万级时不可接受。解决方案：增量索引')
add_bullet('3. LLM API 延迟：Claude API 网络往返约 3-8 秒（取决于生成长度）。这是用户感知延迟的主要来源。解决方案：SSE 流式输出优化首 token 时间（TTFT）体验')
add_bullet('4. 会话存储：内存字典，服务重启后所有会话丢失。解决方案：Redis 持久化')

# ── 4.8 前端 ──
doc.add_heading('4.8 前端', level=2)

doc.add_heading('Q25: 这个 SPA 为什么不用 React/Vue 框架？', level=3)
add_paragraph('这是有意为之的"零构建"（Zero Build Step）设计决策：')
add_bullet('避免 Node.js / npm / webpack 工具链依赖，降低环境配置门槛')
add_bullet('减少部署复杂度：3 个静态文件（index.html + style.css + app.js）通过 FastAPI 直接 serve')
add_bullet('项目定位是个人/小团队工具而非大规模产品，原生 JS 完全够应对当前交互复杂度')
add_bullet('仅通过 CDN 引入 Cytoscape.js 一个外部依赖')
add_bullet('如果未来交互复杂度显著增加（如实时协作），迁移到 Vue/React 也很方便')

doc.add_heading('Q26: SSE 流式在前端怎么处理跨 chunk 的不完整消息？', level=3)
add_paragraph('这是 SSE 客户端实现的核心难点。ReadableStream 的 chunk 边界是任意的（取决于网络 TCP 分段），可能恰好落在 "data: {...}\\n\\n" 的中间。')
add_paragraph('解决方案——buffer 缓冲区：')
add_bullet('每次 read() 后追加到 buffer')
add_bullet('按 \\n 分割 buffer → 最后一行（可能不完整）用 lines.pop() 取出放回 buffer')
add_bullet('完整行才解析 JSON 处理')
add_bullet('下次 read() 时新数据拼接到 buffer 前面，继续尝试解析')
add_paragraph('这是一个经典的流式协议解析模式，等同于 TCP 的"粘包/拆包"处理。')

# ── 4.9 综合问题 ──
doc.add_heading('4.9 综合问题', level=2)

doc.add_heading('Q27: 如果用户问了一个知识库中没有的问题，系统会怎么处理？', level=3)
add_paragraph('当前行为：retriever.retrieve() 仍会返回"最相关"的 chunk（向量检索总是返回最近的 top_k，即使内容实际上不相关）；Agent 无法判断 chunk 是否真正相关，仍会把这些内容传给 LLM；LLM 的 system prompt 指示"上下文不足时明确告知用户"，但这依赖 LLM 自身的判断能力。')
add_paragraph('改进方案：在检索结果上增加相似度阈值过滤。')
add_code('''# 在 agent.py 的 query() 方法中增加阈值判断
if chunks and chunks[0].get("rrf_score", 0) < 0.03:
    return {
        "answer": "No relevant information found in the knowledge base. "
                  "Please try a different question or upload more documents.",
        "sources": [],
        ...
    }''')
add_paragraph('阈值 0.03 需要根据实际数据实验确定（太低会漏掉弱相关但有用的结果，太高会频繁拒绝合理查询）。更好的方案是加入相关性分类器（如用 Cross-encoder 判断 query-chunk 相关性）。')

doc.add_heading('Q28: 这个项目最大的亮点是什么？', level=3)
add_bullet('1. CJK 感知的文本切分器：自研而非依赖 LangChain，解决了中文/日文/韩文文档切分的真实痛点。LangChain 硬编码英文分隔符，完全不处理中文标点，这个项目用约 100 行代码完美解决')
add_bullet('2. 混合检索 + RRF 融合：实现完整且正确，有扎实的理论支撑（信息检索经典论文），不是简单的"调个向量检索 API"')
add_bullet('3. 多 LLM 后端自动检测：用户无需任何配置即可使用，从免费 Ollama 到 Claude API 无缝切换')
add_bullet('4. 全栈能力体现：从文档解析（PDF/DOCX/Web）→ 文本切分 → 向量嵌入 → 混合检索 → LLM Agent → 流式 SSE → 前端可视化，一条完整的 RAG 链路')
add_bullet('5. DOCX XML fallback：展示了解决实际工程兼容性问题的能力，不满足于"库用不了就报错"')
add_bullet('6. 零外部服务依赖：除了 LLM API，所有组件（Embedding、向量数据库、BM25）都在本地运行')

doc.add_heading('Q29: 如果让你重新设计，你会做什么不同的选择？', level=3)
add_bullet('1. 加入 Reranker（Cross-encoder）：在 RRF 之后对 top-20 候选用 Cross-encoder 精排，召回率相近但精度可提升 10-20%（代价是增加约 100ms 延迟）')
add_bullet('2. 异步文档处理：大文档上传后立即返回 HTTP 200，后台用 FastAPI BackgroundTasks 或 Celery 异步处理，避免用户等待')
add_bullet('3. 增量 BM25 索引：避免每次上传都全量重建，改为增量更新 DF 和 corpus')
add_bullet('4. 相似度阈值过滤：在检索结果分数过低时直接返回"未找到"，节省 LLM API 调用费用')
add_bullet('5. Query Expansion（查询改写）：用 LLM 将用户的一个问题改写为多个变体（同义词替换、中英互译），每个变体独立检索后合并结果，提升召回率')
add_bullet('6. 会话持久化：用 Redis 替代内存字典存储会话，支持服务重启不丢失')
add_bullet('7. 可观测性：加入结构化日志（JSON 格式）、Prometheus metrics（检索延迟、LLM 耗时）、用户反馈收集')

doc.add_heading('Q30: 解释项目中用到的设计模式。', level=3)
add_table(
    ['设计模式', '项目中的应用', '说明'],
    [
        ['单例模式', 'config = Config()', '模块级别的全局配置实例，全应用共用'],
        ['策略模式', 'LLMBackend 抽象基类', 'AnthropicBackend / OpenAICompatibleBackend 可互换，客户端代码不感知'],
        ['适配器模式', 'OpenAICompatibleBackend', '将 DeepSeek/Ollama/OpenAI 三个不同 API 适配为统一接口'],
        ['工厂方法', 'LLM._resolve()', '根据环境变量和检测结果创建正确的后端实例'],
        ['懒加载', 'Embedder._model', '延迟加载 80MB 模型，加速启动'],
        ['模板方法', 'DocumentLoader.load()', '统一的加载→清洗→返回 Document 流程，不同格式子方法各自实现'],
        ['管道模式', '_index_document()', 'Loader → Chunker → Embedder → VectorStore → BM25，数据依次流经各阶段'],
        ['责任链模式', 'LLM 自动检测', 'Ollama → Anthropic → DeepSeek → OpenAI 逐级尝试，直到找到可用的'],
    ]
)

# ═══════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════
output_path = r'C:\Users\郝翔\Desktop\AI_Knowledge_Base_Agent_面试准备资料.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
