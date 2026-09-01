"""
Document loader: PDF, DOCX, Markdown, TXT, Web pages.
"""

import re
import uuid
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional

import requests
from bs4 import BeautifulSoup

from ..config import config


@dataclass
class Document:
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:8])
    filename: str = ""
    content: str = ""
    metadata: dict = field(default_factory=dict)
    chunk_ids: List[str] = field(default_factory=list)


class DocumentLoader:
    """Load and parse documents from various formats."""

    @staticmethod
    def load(file_path: str | Path) -> Document:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            content = DocumentLoader._load_pdf(path)
        elif suffix in (".docx", ".doc"):
            content = DocumentLoader._load_docx(path)
        elif suffix in (".md", ".markdown"):
            content = DocumentLoader._load_markdown(path)
        elif suffix in (".txt", ".text", ".csv", ".json", ".xml", ".html", ".htm"):
            content = DocumentLoader._load_text(path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

        content = DocumentLoader._clean_text(content)
        return Document(
            filename=path.name,
            content=content,
            metadata={"source": str(path), "file_type": suffix, "char_count": len(content)},
        )

    @staticmethod
    def load_url(url: str) -> Document:
        headers = {
            "User-Agent": "Mozilla/5.0 (compatible; KnowledgeBot/1.0)"
        }
        resp = requests.get(url, headers=headers, timeout=30)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        text = DocumentLoader._clean_text(text)

        title = ""
        if soup.title and soup.title.string:
            title = soup.title.string.strip()

        return Document(
            filename=title or url,
            content=text,
            metadata={"source": url, "file_type": "web", "char_count": len(text)},
        )

    @staticmethod
    def _load_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            return "\n\n".join(parts)
        except ImportError:
            raise ImportError("pypdf is required for PDF support. pip install pypdf")

    @staticmethod
    def _load_docx(path: Path) -> str:
        """
        从Word文档(.docx)中提取文本内容
        参数:
            path (Path): Word文档的文件路径
        返回:
            str: 提取的文本内容
        异常:
            ImportError: 当缺少python-docx库时抛出
        """
        try:
            # 导入python-docx库
            from docx import Document as DocxDocument
            # 打开Word文档
            doc = DocxDocument(str(path))
            # 提取所有段落文本，忽略空段落
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text += "\n" + row_text

            # Fallback: if python-docx returns nothing, extract directly from XML
            # (handles WPS/older Word formats with text in text boxes)
            if not text.strip():
                text = DocumentLoader._extract_docx_xml(path)

            return text
        except ImportError:
            raise ImportError("python-docx is required. pip install python-docx")

    @staticmethod
    def _extract_docx_xml(path: Path) -> str:
        """Extract text directly from docx XML as fallback, grouped by paragraph."""
        import zipfile
        import re
        paragraphs = []
        with zipfile.ZipFile(str(path)) as zf:
            for name in zf.namelist():
                if name.endswith('.xml') and 'document' in name.lower():
                    xml = zf.read(name).decode('utf-8', errors='ignore')
                    # Group text elements by paragraph (w:p)
                    for para_match in re.finditer(r'<w:p[ >](.*?)</w:p>', xml, re.DOTALL):
                        texts = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', para_match.group(1))
                        line = ''.join(texts).strip()
                        if line:
                            paragraphs.append(line)
        return "\n".join(paragraphs)

    @staticmethod
    def _load_markdown(path: Path) -> str:
        return path.read_text(encoding="utf-8")

    @staticmethod
    def _load_text(path: Path) -> str:
        return path.read_text(encoding="utf-8")

    @staticmethod
    def _clean_text(text: str) -> str:
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]{3,}', '  ', text)
        text = re.sub(r'\x00', '', text)
        return text.strip()
